import os
import uuid
from typing import List, Dict, Any, Tuple
import PyPDF2
from docx import Document
import aiofiles
from datetime import datetime
from app.core.config import settings

class DocumentProcessor:
    def __init__(self):
        self.upload_dir = settings.upload_dir
        os.makedirs(self.upload_dir, exist_ok=True)
    
    async def save_uploaded_file(self, file_content: bytes, filename: str) -> str:
        """Save uploaded file and return the file path."""
        file_id = str(uuid.uuid4())
        file_extension = filename.split('.')[-1].lower()
        safe_filename = f"{file_id}.{file_extension}"
        file_path = os.path.join(self.upload_dir, safe_filename)
        
        async with aiofiles.open(file_path, 'wb') as f:
            await f.write(file_content)
        
        return file_path
    
    def extract_text_from_pdf(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Extract text from PDF file."""
        try:
            text = ""
            metadata = {"page_count": 0, "word_count": 0}
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                metadata["page_count"] = len(pdf_reader.pages)
                
                for page_num, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    text += f"\n--- Page {page_num + 1} ---\n{page_text}\n"
            
            metadata["word_count"] = len(text.split())
            return text, metadata
            
        except Exception as e:
            print(f"Error extracting PDF text: {e}")
            return "", {"page_count": 0, "word_count": 0}
    
    def extract_text_from_docx(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Extract text from DOCX file."""
        try:
            doc = Document(file_path)
            text = ""
            paragraph_count = 0
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
                    paragraph_count += 1
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text += cell.text + " "
                    text += "\n"
            
            metadata = {
                "paragraph_count": paragraph_count,
                "table_count": len(doc.tables),
                "word_count": len(text.split())
            }
            
            return text, metadata
            
        except Exception as e:
            print(f"Error extracting DOCX text: {e}")
            return "", {"paragraph_count": 0, "table_count": 0, "word_count": 0}
    
    def extract_text_from_txt(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Extract text from TXT file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
            
            metadata = {
                "line_count": len(text.split('\n')),
                "word_count": len(text.split()),
                "char_count": len(text)
            }
            
            return text, metadata
            
        except UnicodeDecodeError:
            try:
                with open(file_path, 'r', encoding='latin-1') as file:
                    text = file.read()
                metadata = {
                    "line_count": len(text.split('\n')),
                    "word_count": len(text.split()),
                    "char_count": len(text)
                }
                return text, metadata
            except Exception as e:
                print(f"Error extracting TXT text: {e}")
                return "", {"line_count": 0, "word_count": 0, "char_count": 0}
        except Exception as e:
            print(f"Error extracting TXT text: {e}")
            return "", {"line_count": 0, "word_count": 0, "char_count": 0}
    
    def extract_text_from_file(self, file_path: str, filename: str) -> Tuple[str, Dict[str, Any]]:
        """Extract text from any supported file type."""
        file_extension = filename.split('.')[-1].lower()
        
        if file_extension == 'pdf':
            return self.extract_text_from_pdf(file_path)
        elif file_extension == 'docx':
            return self.extract_text_from_docx(file_path)
        elif file_extension in ['txt', 'md']:
            return self.extract_text_from_txt(file_path)
        else:
            return "", {"error": f"Unsupported file type: {file_extension}"}
    
    def chunk_text(self, text: str, filename: str = "", chunk_size: int = None, chunk_overlap: int = None) -> List[Dict[str, Any]]:
        """Split text into chunks for vector storage."""
        if chunk_size is None:
            chunk_size = settings.chunk_size
        if chunk_overlap is None:
            chunk_overlap = settings.chunk_overlap
        
        chunks = []
        text_length = len(text)
        
        if text_length <= chunk_size:
            chunks.append({
                "content": text,
                "chunk_index": 0,
                "start_char": 0,
                "end_char": text_length,
                "metadata": {
                    "filename": filename,
                    "total_chunks": 1
                }
            })
            return chunks
        
        start_pos = 0
        chunk_index = 0
        
        while start_pos < text_length:
            end_pos = min(start_pos + chunk_size, text_length)
            
            # Try to break at sentence boundaries
            if end_pos < text_length:
                # Look for sentence endings within the last 100 characters
                search_start = max(end_pos - 100, start_pos)
                sentence_endings = ['.', '!', '?', '\n\n']
                
                best_break = -1
                for ending in sentence_endings:
                    pos = text.rfind(ending, search_start, end_pos)
                    if pos > best_break:
                        best_break = pos
                
                if best_break > start_pos:
                    end_pos = best_break + 1
            
            chunk_content = text[start_pos:end_pos].strip()
            
            if chunk_content:
                chunks.append({
                    "content": chunk_content,
                    "chunk_index": chunk_index,
                    "start_char": start_pos,
                    "end_char": end_pos,
                    "metadata": {
                        "filename": filename,
                        "chunk_size": len(chunk_content)
                    }
                })
                chunk_index += 1
            
            # Move start position, accounting for overlap
            start_pos = max(start_pos + chunk_size - chunk_overlap, end_pos)
        
        # Update total chunks in metadata
        for chunk in chunks:
            chunk["metadata"]["total_chunks"] = len(chunks)
        
        return chunks
    
    def delete_file(self, file_path: str) -> bool:
        """Delete a file from storage."""
        try:
            if os.path.exists(file_path):
                os.remove(file_path)
                return True
            return False
        except Exception as e:
            print(f"Error deleting file: {e}")
            return False
    
    def get_file_info(self, file_path: str) -> Dict[str, Any]:
        """Get basic file information."""
        try:
            if not os.path.exists(file_path):
                return {"error": "File not found"}
            
            stat = os.stat(file_path)
            return {
                "size": stat.st_size,
                "created": datetime.fromtimestamp(stat.st_ctime),
                "modified": datetime.fromtimestamp(stat.st_mtime),
                "exists": True
            }
        except Exception as e:
            return {"error": str(e), "exists": False}

# Global instance
document_processor = DocumentProcessor() 